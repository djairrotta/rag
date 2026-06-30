"""Renderiza o texto do recurso em DOCX (python-docx).

Em produção, o DOCX pode ser convertido a PDF via LibreOffice headless (blueprint);
aqui geramos o DOCX, que já é entregável e editável.
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def render_docx(title: str, body: str) -> bytes:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    doc.add_paragraph("")

    for line in (body or "").replace("\r\n", "\n").split("\n"):
        line = line.rstrip()
        p = doc.add_paragraph(line)
        # títulos de seção em negrito
        upper = line.strip()
        if upper[:3] in ("I —", "II ", "III") or upper.startswith(("I —", "II —", "III —")):
            for r in p.runs:
                r.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
